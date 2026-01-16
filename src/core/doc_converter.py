#!/usr/bin/env python3
"""
Document Converter for TERAG AI-REPS
Конвертирует различные форматы документов в текстовые файлы для анализа
"""

import os
import sys
from pathlib import Path
from typing import Dict, List, Optional
import logging
from langdetect import detect

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def detect_language(text: str) -> str:
    """Определяет язык текста"""
    try:
        if len(text.strip()) < 10:
            return "unknown"
        lang = detect(text)
        return "ru" if lang == "ru" else "en"
    except Exception as e:
        logger.warning(f"Language detection failed: {e}")
        return "unknown"

def extract_text_from_docx(path: Path) -> str:
    """Извлечение текста из DOCX файлов"""
    try:
        from docx import Document
        doc = Document(path)
        text_parts = []
        
        # Извлекаем текст из параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Ошибка извлечения текста из DOCX {path}: {e}")
        return ""

def extract_text_from_pdf(path: Path) -> str:
    """Извлечение текста из PDF файлов"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(path)
        text_parts = []
        
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(f"--- Страница {page_num + 1} ---\n{page_text.strip()}")
            except Exception as e:
                logger.warning(f"Ошибка извлечения текста со страницы {page_num + 1} в {path}: {e}")
                continue
        
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"Ошибка извлечения текста из PDF {path}: {e}")
        return ""

def extract_text_from_xlsx(path: Path) -> str:
    """Извлечение текста из XLSX файлов"""
    try:
        import pandas as pd
        text_parts = []
        
        # Читаем все листы
        excel_file = pd.ExcelFile(path)
        for sheet_name in excel_file.sheet_names:
            try:
                df = pd.read_excel(path, sheet_name=sheet_name)
                if not df.empty:
                    text_parts.append(f"--- Лист: {sheet_name} ---")
                    text_parts.append(df.to_string(index=False, na_rep=''))
            except Exception as e:
                logger.warning(f"Ошибка чтения листа {sheet_name} в {path}: {e}")
                continue
        
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"Ошибка извлечения текста из XLSX {path}: {e}")
        return ""

def extract_text_from_txt(path: Path) -> str:
    """Чтение текстовых файлов"""
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(path, 'r', encoding='cp1251') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Ошибка чтения текстового файла {path}: {e}")
            return ""

def convert_all_to_txt(data_dir: str = "data") -> Path:
    """Конвертация всех документов в текстовые файлы"""
    data_path = Path(data_dir)
    txt_dir = data_path / "converted"
    txt_dir.mkdir(parents=True, exist_ok=True)
    
    # Маппинг расширений на функции извлечения
    extractors = {
        ".docx": extract_text_from_docx,
        ".pdf": extract_text_from_pdf,
        ".xlsx": extract_text_from_xlsx,
        ".xls": extract_text_from_xlsx,  # Старый формат Excel
        ".txt": extract_text_from_txt,
        ".md": extract_text_from_txt,
    }
    
    converted_count = 0
    error_count = 0
    
    logger.info(f"🔍 Поиск документов в {data_path}")
    
    # Рекурсивно ищем все файлы
    for ext, extractor in extractors.items():
        files = list(data_path.rglob(f"*{ext}"))
        logger.info(f"📄 Найдено {len(files)} файлов с расширением {ext}")
        
        for file_path in files:
            try:
                # Пропускаем уже конвертированные файлы
                if "converted" in str(file_path):
                    continue
                
                logger.info(f"🔄 Обрабатываем: {file_path.name}")
                text = extractor(file_path)
                
                if text.strip():
                    # Определяем язык текста
                    language = detect_language(text)
                    
                    # Создаём имя выходного файла с указанием языка
                    output_name = f"{file_path.stem}_{language}.txt"
                    output_path = txt_dir / output_name
                    
                    # Сохраняем текст
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(f"# Исходный файл: {file_path.name}\n")
                        f.write(f"# Путь: {file_path}\n")
                        f.write(f"# Размер: {len(text)} символов\n")
                        f.write(f"# Язык: {language}\n\n")
                        f.write(text)
                    
                    converted_count += 1
                    logger.info(f"✅ Конвертирован: {file_path.name} → {output_name} ({len(text)} символов, язык: {language})")
                else:
                    logger.warning(f"⚠️ Пустой файл: {file_path.name}")
                    
            except Exception as e:
                error_count += 1
                logger.error(f"❌ Ошибка при обработке {file_path.name}: {e}")
    
    logger.info(f"📊 Конвертация завершена:")
    logger.info(f"   ✅ Успешно: {converted_count} файлов")
    logger.info(f"   ❌ Ошибок: {error_count} файлов")
    logger.info(f"   📁 Результат: {txt_dir}")
    
    return txt_dir

def get_converted_files(converted_dir: Path) -> List[Path]:
    """Получение списка конвертированных файлов"""
    if not converted_dir.exists():
        return []
    
    txt_files = list(converted_dir.glob("*.txt"))
    logger.info(f"📚 Найдено {len(txt_files)} конвертированных файлов")
    return txt_files

def analyze_document_structure(converted_dir: Path) -> Dict[str, any]:
    """Анализ структуры документов"""
    txt_files = get_converted_files(converted_dir)
    
    if not txt_files:
        return {"error": "Нет конвертированных файлов"}
    
    analysis = {
        "total_files": len(txt_files),
        "total_size": 0,
        "file_sizes": [],
        "sources": set(),
        "years": set(),
        "file_types": {}
    }
    
    for file_path in txt_files:
        try:
            # Размер файла
            file_size = file_path.stat().st_size
            analysis["total_size"] += file_size
            analysis["file_sizes"].append(file_size)
            
            # Читаем заголовок для анализа
            with open(file_path, 'r', encoding='utf-8') as f:
                header = f.read(500)  # Первые 500 символов
            
            # Извлекаем информацию из заголовка
            if "Исходный файл:" in header:
                source_line = [line for line in header.split('\n') if "Исходный файл:" in line]
                if source_line:
                    source_file = source_line[0].split(": ")[1]
                    analysis["sources"].add(source_file)
                    
                    # Пытаемся извлечь год из имени файла
                    import re
                    year_match = re.search(r'20\d{2}', source_file)
                    if year_match:
                        analysis["years"].add(year_match.group())
            
            # Определяем тип по расширению исходного файла
            if ".docx" in header:
                analysis["file_types"]["docx"] = analysis["file_types"].get("docx", 0) + 1
            elif ".pdf" in header:
                analysis["file_types"]["pdf"] = analysis["file_types"].get("pdf", 0) + 1
            elif ".xlsx" in header or ".xls" in header:
                analysis["file_types"]["excel"] = analysis["file_types"].get("excel", 0) + 1
            else:
                analysis["file_types"]["other"] = analysis["file_types"].get("other", 0) + 1
                
        except Exception as e:
            logger.warning(f"Ошибка анализа файла {file_path}: {e}")
    
    # Конвертируем set в list для JSON сериализации
    analysis["sources"] = list(analysis["sources"])
    analysis["years"] = list(analysis["years"])
    
    return analysis

if __name__ == "__main__":
    # Запуск конвертации
    converted_dir = convert_all_to_txt()
    
    # Анализ результатов
    analysis = analyze_document_structure(converted_dir)
    
    print("\n" + "="*50)
    print("📊 АНАЛИЗ ДОКУМЕНТОВ")
    print("="*50)
    print(f"📁 Всего файлов: {analysis.get('total_files', 0)}")
    print(f"📏 Общий размер: {analysis.get('total_size', 0) / 1024:.1f} KB")
    print(f"📅 Годы: {', '.join(sorted(analysis.get('years', [])))}")
    print(f"📄 Типы файлов: {analysis.get('file_types', {})}")
    print(f"📚 Источники: {len(analysis.get('sources', []))} файлов")
    print("="*50)
